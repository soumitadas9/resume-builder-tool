import os
import argparse
import json
import re
from docx import Document
from docx2txt import process
from groq import Groq
from PyPDF2 import PdfReader

SKILL_KEYWORDS = {
    'python', 'pytorch', 'tensorflow', 'numpy', 'pandas', 'scikit-learn', 'scikit', 'spark', 'hadoop', 'sql',
    'postgresql', 'postgres', 'mysql', 'mongodb', 'redis', 'aws', 'azure', 'gcp', 'google cloud', 'docker',
    'kubernetes', 'git', 'github', 'gitlab', 'linux', 'tableau', 'power bi', 'look', 'looker', 'apache',
    'matplotlib', 'seaborn', 'keras', 'statsmodels', 'jupyter', 'streamlit', 'flask', 'django', 'airflow',
    'hdfs', 'mlflow', 'salesforce', 'snowflake', 'dbt', 'databricks', 'bigquery', 'sagemaker',
    'bash', 'shell', 'unix', 'ci/cd', 'etl', 'rest', 'api', 'json', 'graph', 'sqlalchemy',
    'xgboost', 'lightgbm', 'catboost', 'opencv', 'spacy', 'nltk', 'gensim', 'pytorch-lightning',
    'excel', 'word', 'powerpoint', 'lookml'
}

SKILL_PHRASES = {
    'machine learning', 'deep learning', 'natural language processing', 'computer vision', 'data engineering',
    'project management', 'software engineering', 'quality assurance', 'cloud computing', 'business intelligence',
    'risk management', 'data science', 'data analytics', 'information security', 'cyber security', 'visualization',
    'feature engineering', 'model deployment', 'a/b testing', 'agile', 'scrum', 'kanban',
    'product management', 'time management', 'critical thinking', 'problem solving',
    'communication', 'teamwork', 'leadership', 'collaboration'
}


def parse_llm_json_response(text):
    """Parse JSON from LLM output, even if wrapped in extraneous text."""
    cleaned = re.sub(r'```(?:json)?\s*', '', text)
    cleaned = re.sub(r'```', '', cleaned)

    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass

    # Attempt a safer parse for JSON-like output with raw newlines inside strings.
    missing_skills = []
    roadmap = None

    skills_match = re.search(r'"missing_skills"\s*:\s*\[([\s\S]*?)\]', cleaned)
    if skills_match:
        list_text = skills_match.group(1)
        missing_skills = re.findall(r'"([^"]+)"', list_text)

    roadmap_match = re.search(r'"roadmap"\s*:\s*"([\s\S]*?)"\s*(?:,|\})', cleaned)
    if roadmap_match:
        roadmap_text = roadmap_match.group(1)
        roadmap_text = roadmap_text.strip()
        roadmap_text = roadmap_text.replace('\\"', '"')
        roadmap = roadmap_text

    if missing_skills or roadmap is not None:
        return {
            'missing_skills': missing_skills,
            'roadmap': roadmap
        }

    return None


def local_extract_skills(text):
    """Extract likely skills from text for local fallback."""
    lower_text = text.lower()
    skills = set()
    for phrase in SKILL_PHRASES:
        if phrase in lower_text:
            skills.add(phrase)
    for keyword in SKILL_KEYWORDS:
        if keyword in lower_text:
            skills.add(keyword)
    return skills


def llm_generate_missing_skills_and_roadmap(resume, job_description):
    """Use the LLM to generate missing skills and a roadmap."""
    groq_api_key = os.getenv('GROQ_API_KEY')
    if not groq_api_key:
        raise RuntimeError(
            'GROQ_API_KEY is not set. Please export your Groq API key before running the script.'
        )

    client = Groq(api_key=groq_api_key)
    messages = [
        {
            "role": "system",
            "content": (
                "You are an expert resume assistant. Output only valid JSON. "
                "Do not include markdown, explanations, or any text outside the JSON object. "
                "If you cannot identify missing skills, return an empty list for `missing_skills`. "
                "If you cannot create a roadmap, return an empty string for `roadmap`."
            )
        },
        {
            "role": "user",
            "content": (
                "Identify skills required by the job description that are missing from the resume. "
                "Provide a concise, actionable roadmap for acquiring or demonstrating those skills. "
                "Return only valid JSON with exactly two keys: `missing_skills` and `roadmap`. "
                "`missing_skills` must be a JSON array of distinct strings. "
                "`roadmap` must be a string containing concrete steps and resources.\n\n"
                "Resume:\n" + resume + "\n\n"
                "Job description:\n" + job_description
            )
        }
    ]

    completion = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=messages,
        temperature=0.2,
        max_completion_tokens=1024
    )

    def extract_text(chunk):
        if chunk is None:
            return ''
        if isinstance(chunk, str):
            return chunk
        if isinstance(chunk, bytes):
            return chunk.decode('utf-8', errors='ignore')
        if isinstance(chunk, (list, tuple)):
            return ''.join(extract_text(item) for item in chunk)

        if hasattr(chunk, 'choices'):
            choices = getattr(chunk, 'choices')
            if not choices:
                return ''
            choice = choices[0]
            if hasattr(choice, 'delta'):
                return getattr(choice.delta, 'content', '') or ''
            if hasattr(choice, 'message'):
                return getattr(choice.message, 'content', '') or ''

        if isinstance(chunk, dict):
            choices = chunk.get('choices')
            if choices:
                choice = choices[0]
                delta = choice.get('delta')
                if delta:
                    return delta.get('content', '') or ''
                message = choice.get('message')
                if message:
                    return message.get('content', '') or ''

        return ''

    output = ''
    if hasattr(completion, 'choices') or isinstance(completion, dict):
        output = extract_text(completion)
    else:
        for chunk in completion:
            output += extract_text(chunk)

    data = parse_llm_json_response(output)
    if not data:
        with open('llm_raw_output.txt', 'w', encoding='utf-8') as raw_file:
            raw_file.write(output)
        return None, None

    missing_skills = data.get('missing_skills')
    roadmap = data.get('roadmap')
    if isinstance(missing_skills, str):
        missing_skills = [item.strip() for item in missing_skills.split(',') if item.strip()]

    return missing_skills, roadmap


def parse_pdf_file(filename):
    """Parse a PDF file to plain text."""
    text_parts = []
    reader = PdfReader(filename)
    for page in reader.pages:
        text = page.extract_text() or ''
        if text:
            text_parts.append(text)
    return '\n'.join(text_parts)


def parse_doc_file(filename):
    """Parse a DOC file to plain text."""
    return process(filename) or ''


def parse_docx_file(filename):
    """Parse a DOCX file to plain text."""
    doc = Document(filename)
    return ' '.join(para.text for para in doc.paragraphs)


def parse_txt_file(filename):
    """Parse a TXT file to plain text."""
    with open(filename, 'r', encoding='utf-8') as f:
        return f.read()


def get_page_size_from_pdf(pdf_path):
    """Return the width and height of the first page in the PDF."""
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError('PyMuPDF is required for PDF page size support.') from exc

    with fitz.open(pdf_path) as doc:
        page = doc[0]
        rect = page.rect
        return rect.width, rect.height


def wrap_text(text, font, fontsize, max_width):
    """Wrap text into lines that fit within max_width."""
    words = text.split()
    lines = []
    line = ''
    for word in words:
        candidate = word if not line else f"{line} {word}"
        if font.text_length(candidate, fontsize=fontsize) <= max_width:
            line = candidate
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines


def save_text_as_pdf(text, output_path, template_pdf=None):
    """Save text to a PDF; if template_pdf is provided, use its page size."""
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError('PyMuPDF is required to save text as PDF.') from exc

    if template_pdf and os.path.exists(template_pdf):
        page_width, page_height = get_page_size_from_pdf(template_pdf)
    else:
        page_width, page_height = 595, 842  # A4 points

    doc = fitz.open()
    font = fitz.Font()
    margin = 48
    line_height = 16
    max_y = page_height - margin
    current_y = margin
    usable_width = page_width - margin * 2
    page = doc.new_page(width=page_width, height=page_height)

    for paragraph in text.split('\n'):
        if not paragraph.strip():
            current_y += line_height
            if current_y + line_height > max_y:
                page = doc.new_page(width=page_width, height=page_height)
                current_y = margin
            continue

        lines = wrap_text(paragraph, font, 12, usable_width)
        for line in lines:
            if current_y + line_height > max_y:
                page = doc.new_page(width=page_width, height=page_height)
                current_y = margin
            page.insert_text((margin, current_y), line, fontsize=12, fontname="helv")
            current_y += line_height

        current_y += line_height
        if current_y + line_height > max_y:
            page = doc.new_page(width=page_width, height=page_height)
            current_y = margin

    doc.save(output_path)


def save_text_as_docx(text, output_path, template_docx=None):
    """Save text to a DOCX file, reusing a template document if available."""
    if template_docx and os.path.exists(template_docx):
        doc = Document(template_docx)
    else:
        doc = Document()

    lines = text.split('\n')
    paragraph_count = len(doc.paragraphs)
    idx = 0

    for line in lines:
        if idx < paragraph_count:
            doc.paragraphs[idx].text = line
        else:
            doc.add_paragraph(line)
        idx += 1

    for remaining in range(idx, paragraph_count):
        doc.paragraphs[remaining].text = ''

    doc.save(output_path)


def gen_local_roadmap(missing_skills):
    if not missing_skills:
        return (
            "Your resume already aligns well with the job description. Continue strengthening your existing skills, "
            "build more project examples, and update your resume regularly with measurable accomplishments."
        )

    lines = [
        "Recommended roadmap to acquire and demonstrate missing skills:",
        ""
    ]
    for skill in missing_skills:
        lines.append(f"- {skill}: Learn through targeted online courses, build a hands-on project, and add concrete examples to your resume.")
    lines.extend([
        "",
        "General steps:",
        "1. Identify the highest-priority missing skills and select one or two to focus on first.",
        "2. Complete a practical course or tutorial that includes a project for each skill.",
        "3. Build or update a small portfolio project that demonstrates the skill in context.",
        "4. Document the work as bullet points in your resume, emphasizing results and tools used.",
        "5. Review the job description again and tailor the resume with those skill examples."
    ])
    return '\n'.join(lines)


def infer_input_format(input_file):
    ext = os.path.splitext(input_file)[1].lower()
    if ext == '.pdf':
        return 'pdf'
    if ext == '.doc':
        return 'doc'
    if ext == '.docx':
        return 'docx'
    if ext == '.txt':
        return 'txt'
    raise ValueError(f"Cannot infer input format from extension '{ext}'. Please use -f/--input_format.")


def main(args):
    if not os.path.exists(args.input_file):
        raise FileNotFoundError(f"Input file not found: {args.input_file}")
    if not os.path.exists(args.job_description):
        raise FileNotFoundError(f"Job description file not found: {args.job_description}")

    input_format = args.input_format or infer_input_format(args.input_file)

    if input_format == 'pdf':
        input_text = parse_pdf_file(args.input_file)
    elif input_format == 'doc':
        input_text = parse_doc_file(args.input_file)
    elif input_format == 'docx':
        input_text = parse_docx_file(args.input_file)
    elif input_format == 'txt':
        input_text = parse_txt_file(args.input_file)
    else:
        raise ValueError(f"Unsupported input format: {input_format}")

    job_description_text = parse_txt_file(args.job_description)

    llm_missing_skills, llm_roadmap_text = llm_generate_missing_skills_and_roadmap(input_text, job_description_text)
    if llm_missing_skills is None or not isinstance(llm_missing_skills, list):
        print('ERROR: LLM response could not be parsed. Using local fallback extraction to generate missing skills and roadmap.')
        resume_skills = local_extract_skills(input_text)
        job_skills = local_extract_skills(job_description_text)
        missing_skills = sorted(job_skills - resume_skills)
        roadmap_text = gen_local_roadmap(missing_skills)
    else:
        print('SUCCESS: LLM output parsed successfully. Generating missing_skills.txt and roadmap.txt from LLM response.')
        missing_skills = [skill.strip() for skill in llm_missing_skills if skill.strip()]
        roadmap_text = llm_roadmap_text.strip() if isinstance(llm_roadmap_text, str) and llm_roadmap_text.strip() else gen_local_roadmap(missing_skills)

    with open('missing_skills.txt', 'w', encoding='utf-8') as f:
        for skill in sorted(missing_skills):
            f.write(f"{skill}\n")

    with open('roadmap.txt', 'w', encoding='utf-8') as f:
        f.write(roadmap_text)

    print('Saved missing skills to missing_skills.txt')
    print('Saved roadmap to roadmap.txt')


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Find missing skills and generate a learning roadmap')
    parser.add_argument('-i', '--input_file', help='Input file (PDF, DOC, DOCX, or TXT)', required=True)
    parser.add_argument('-f', '--input_format', choices=['pdf', 'doc', 'docx', 'txt'], help='Format of the input resume file (optional; inferred from file extension if omitted)')
    parser.add_argument('-j', '--job_description', help='Job description file (TXT)', required=True)
    args = parser.parse_args()
    main(args)
